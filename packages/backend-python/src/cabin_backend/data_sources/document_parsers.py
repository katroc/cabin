"""
Document parsers for extracting text and metadata from various file formats.
Supports PDF, DOCX, Markdown, HTML, and plain text files with comprehensive metadata extraction.
"""

import logging
import mimetypes
import os
import re
from abc import ABC, abstractmethod
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, Optional, List, Tuple

logger = logging.getLogger(__name__)


class DocumentMetadata:
    """Container for document metadata extracted from files."""

    def __init__(self):
        # File properties
        self.filename: Optional[str] = None
        self.file_size: Optional[int] = None
        self.file_extension: Optional[str] = None
        self.mime_type: Optional[str] = None
        self.created_at: Optional[datetime] = None
        self.modified_at: Optional[datetime] = None

        # Document properties
        self.title: Optional[str] = None
        self.author: Optional[str] = None
        self.subject: Optional[str] = None
        self.keywords: List[str] = []
        self.language: Optional[str] = None
        self.page_count: Optional[int] = None
        self.word_count: Optional[int] = None
        self.character_count: Optional[int] = None

        # Content analysis
        self.has_images: bool = False
        self.has_tables: bool = False
        self.heading_count: int = 0
        self.headings: List[str] = []

        # Processing metadata
        self.parser_used: Optional[str] = None
        self.extraction_warnings: List[str] = []
        self.is_encrypted: bool = False
        self.is_corrupted: bool = False

    def to_dict(self) -> Dict[str, Any]:
        """Convert metadata to dictionary for storage."""
        return {
            "filename": self.filename,
            "file_size": self.file_size,
            "file_extension": self.file_extension,
            "mime_type": self.mime_type,
            "created_at": self.created_at.isoformat() if self.created_at else None,
            "modified_at": self.modified_at.isoformat() if self.modified_at else None,
            "title": self.title,
            "page_title": self.title or self.filename,  # Fallback for metadata enrichment
            "author": self.author,
            "subject": self.subject,
            "keywords": self.keywords,
            "language": self.language,
            "page_count": self.page_count,
            "word_count": self.word_count,
            "character_count": self.character_count,
            "has_images": self.has_images,
            "has_tables": self.has_tables,
            "heading_count": self.heading_count,
            "headings": self.headings,
            "parser_used": self.parser_used,
            "extraction_warnings": self.extraction_warnings,
            "is_encrypted": self.is_encrypted,
            "is_corrupted": self.is_corrupted,
        }


class DocumentParser(ABC):
    """Abstract base class for document parsers."""

    @abstractmethod
    def can_parse(self, file_path: Path) -> bool:
        """Check if this parser can handle the given file."""
        pass

    @abstractmethod
    def parse(self, file_path: Path) -> Tuple[str, DocumentMetadata]:
        """
        Parse document and extract text and metadata.
        Returns tuple of (extracted_text, metadata).
        """
        pass

    def _get_file_metadata(self, file_path: Path) -> DocumentMetadata:
        """Extract basic file system metadata."""
        metadata = DocumentMetadata()

        try:
            stat = file_path.stat()
            metadata.filename = file_path.name
            metadata.file_size = stat.st_size
            metadata.file_extension = file_path.suffix.lower()
            metadata.created_at = datetime.fromtimestamp(stat.st_ctime)
            metadata.modified_at = datetime.fromtimestamp(stat.st_mtime)

            # Detect MIME type
            mime_type, _ = mimetypes.guess_type(str(file_path))
            metadata.mime_type = mime_type

        except Exception as e:
            logger.warning(f"Failed to extract file metadata for {file_path}: {e}")
            metadata.extraction_warnings.append(f"File metadata extraction failed: {e}")

        return metadata

    def _analyze_text_content(self, text: str, metadata: DocumentMetadata) -> None:
        """Analyze text content and update metadata."""
        if not text:
            return

        # Basic text statistics
        metadata.character_count = len(text)
        metadata.word_count = len(text.split())

        # Extract headings (markdown-style or HTML)
        heading_patterns = [
            r'^#{1,6}\s+(.+)$',  # Markdown headings
            r'<h[1-6][^>]*>([^<]+)</h[1-6]>',  # HTML headings
        ]

        headings = []
        for pattern in heading_patterns:
            matches = re.findall(pattern, text, re.MULTILINE | re.IGNORECASE)
            headings.extend(matches)

        metadata.headings = list(set(headings))  # Remove duplicates
        metadata.heading_count = len(metadata.headings)

        # Check for tables and images
        metadata.has_tables = bool(re.search(r'<table|^\|.*\|', text, re.MULTILINE | re.IGNORECASE))
        metadata.has_images = bool(re.search(r'<img|!\[.*\]\(', text, re.IGNORECASE))


class PDFParser(DocumentParser):
    """Parser for PDF documents using PyPDF2."""

    def can_parse(self, file_path: Path) -> bool:
        return file_path.suffix.lower() == '.pdf'

    def parse(self, file_path: Path) -> Tuple[str, DocumentMetadata]:
        metadata = self._get_file_metadata(file_path)
        metadata.parser_used = "PDFParser"

        try:
            import PyPDF2
        except ImportError:
            logger.error("PyPDF2 not available for PDF parsing")
            metadata.extraction_warnings.append("PyPDF2 not available")
            return "", metadata

        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)

                # Check for encryption
                if reader.is_encrypted:
                    metadata.is_encrypted = True
                    metadata.extraction_warnings.append("PDF is encrypted")
                    return "", metadata

                # Extract metadata from PDF
                if reader.metadata:
                    pdf_meta = reader.metadata
                    metadata.title = pdf_meta.get('/Title')
                    metadata.author = pdf_meta.get('/Author')
                    metadata.subject = pdf_meta.get('/Subject')

                    # Parse keywords
                    keywords = pdf_meta.get('/Keywords', '')
                    if keywords:
                        metadata.keywords = [k.strip() for k in keywords.split(',')]

                # Fallback: use title or filename for page_title
                if not metadata.title and metadata.filename:
                    metadata.title = metadata.filename

                # Extract text from all pages
                metadata.page_count = len(reader.pages)
                text_content = []

                for page_num, page in enumerate(reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content.append(page_text)
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                        metadata.extraction_warnings.append(f"Page {page_num + 1} extraction failed")

                full_text = '\n\n'.join(text_content)
                self._analyze_text_content(full_text, metadata)

                return full_text, metadata

        except Exception as e:
            logger.error(f"Failed to parse PDF {file_path}: {e}")
            metadata.is_corrupted = True
            metadata.extraction_warnings.append(f"PDF parsing failed: {e}")
            return "", metadata


class DOCXParser(DocumentParser):
    """Parser for DOCX documents using python-docx."""

    def can_parse(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in ['.docx', '.docm']

    def parse(self, file_path: Path) -> Tuple[str, DocumentMetadata]:
        metadata = self._get_file_metadata(file_path)
        metadata.parser_used = "DOCXParser"

        try:
            from docx import Document
        except ImportError:
            logger.error("python-docx not available for DOCX parsing")
            metadata.extraction_warnings.append("python-docx not available")
            return "", metadata

        try:
            doc = Document(file_path)

            # Extract document properties
            core_props = doc.core_properties
            metadata.title = core_props.title
            metadata.author = core_props.author
            metadata.subject = core_props.subject
            metadata.keywords = core_props.keywords.split(',') if core_props.keywords else []
            metadata.language = core_props.language
            metadata.created_at = core_props.created
            metadata.modified_at = core_props.modified

            # Extract text content
            text_content = []
            heading_count = 0

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    # Check if it's a heading
                    if paragraph.style.name.startswith('Heading'):
                        heading_count += 1
                        metadata.headings.append(paragraph.text.strip())

                    text_content.append(paragraph.text)

            # Check for tables
            if doc.tables:
                metadata.has_tables = True
                for table in doc.tables:
                    table_text = []
                    for row in table.rows:
                        row_text = [cell.text.strip() for cell in row.cells]
                        table_text.append(' | '.join(row_text))
                    text_content.append('\n'.join(table_text))

            # Check for images
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    metadata.has_images = True
                    break

            full_text = '\n\n'.join(text_content)
            metadata.heading_count = heading_count
            self._analyze_text_content(full_text, metadata)

            return full_text, metadata

        except Exception as e:
            logger.error(f"Failed to parse DOCX {file_path}: {e}")
            metadata.is_corrupted = True
            metadata.extraction_warnings.append(f"DOCX parsing failed: {e}")
            return "", metadata


class MarkdownParser(DocumentParser):
    """Parser for Markdown documents."""

    def can_parse(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in ['.md', '.markdown', '.mdown', '.mkd']

    def parse(self, file_path: Path) -> Tuple[str, DocumentMetadata]:
        metadata = self._get_file_metadata(file_path)
        metadata.parser_used = "MarkdownParser"

        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()

            # Extract frontmatter metadata if present
            if content.startswith('---'):
                try:
                    import yaml
                    parts = content.split('---', 2)
                    if len(parts) >= 3:
                        frontmatter = yaml.safe_load(parts[1])
                        content = parts[2].strip()

                        # Extract common frontmatter fields
                        if isinstance(frontmatter, dict):
                            metadata.title = frontmatter.get('title')
                            metadata.author = frontmatter.get('author')
                            metadata.subject = frontmatter.get('subject')
                            metadata.keywords = frontmatter.get('keywords', [])
                            if isinstance(metadata.keywords, str):
                                metadata.keywords = [k.strip() for k in metadata.keywords.split(',')]
                except Exception as e:
                    logger.warning(f"Failed to parse frontmatter: {e}")
                    metadata.extraction_warnings.append("Frontmatter parsing failed")

            # If no title from frontmatter, try to extract from first heading
            if not metadata.title:
                first_heading = re.search(r'^#\s+(.+)$', content, re.MULTILINE)
                if first_heading:
                    metadata.title = first_heading.group(1).strip()

            self._analyze_text_content(content, metadata)
            return content, metadata

        except Exception as e:
            logger.error(f"Failed to parse Markdown {file_path}: {e}")
            metadata.is_corrupted = True
            metadata.extraction_warnings.append(f"Markdown parsing failed: {e}")
            return "", metadata


class HTMLParser(DocumentParser):
    """Parser for HTML documents."""

    def can_parse(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in ['.html', '.htm']

    def parse(self, file_path: Path) -> Tuple[str, DocumentMetadata]:
        metadata = self._get_file_metadata(file_path)
        metadata.parser_used = "HTMLParser"

        try:
            from bs4 import BeautifulSoup
        except ImportError:
            logger.error("BeautifulSoup not available for HTML parsing")
            metadata.extraction_warnings.append("BeautifulSoup not available")
            return "", metadata

        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()

            soup = BeautifulSoup(content, 'html.parser')

            # Extract metadata from HTML head
            title_tag = soup.find('title')
            if title_tag:
                metadata.title = title_tag.get_text().strip()

            # Extract meta tags
            meta_tags = soup.find_all('meta')
            for meta in meta_tags:
                name = meta.get('name', '').lower()
                content_attr = meta.get('content', '')

                if name == 'author':
                    metadata.author = content_attr
                elif name == 'description':
                    metadata.subject = content_attr
                elif name == 'keywords':
                    metadata.keywords = [k.strip() for k in content_attr.split(',')]
                elif name == 'language':
                    metadata.language = content_attr

            # Check for images and tables
            metadata.has_images = bool(soup.find('img'))
            metadata.has_tables = bool(soup.find('table'))

            # Extract headings
            headings = soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6'])
            metadata.headings = [h.get_text().strip() for h in headings]
            metadata.heading_count = len(metadata.headings)

            # Extract text content
            text_content = soup.get_text(separator=' ', strip=True)
            self._analyze_text_content(text_content, metadata)

            return content, metadata  # Return original HTML for processing

        except Exception as e:
            logger.error(f"Failed to parse HTML {file_path}: {e}")
            metadata.is_corrupted = True
            metadata.extraction_warnings.append(f"HTML parsing failed: {e}")
            return "", metadata


class TextParser(DocumentParser):
    """Parser for plain text documents."""

    def can_parse(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in ['.txt', '.text', '.log', '.csv']

    def parse(self, file_path: Path) -> Tuple[str, DocumentMetadata]:
        metadata = self._get_file_metadata(file_path)
        metadata.parser_used = "TextParser"

        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            content = None

            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                    break
                except UnicodeDecodeError:
                    continue

            if content is None:
                raise ValueError("Could not decode file with any supported encoding")

            # Try to extract title from filename or first line
            if not metadata.title:
                first_line = content.split('\n')[0].strip()
                if len(first_line) < 100 and first_line:
                    metadata.title = first_line
                else:
                    metadata.title = file_path.stem

            # Check for CSV structure
            if file_path.suffix.lower() == '.csv':
                metadata.has_tables = True

            self._analyze_text_content(content, metadata)
            return content, metadata

        except Exception as e:
            logger.error(f"Failed to parse text file {file_path}: {e}")
            metadata.is_corrupted = True
            metadata.extraction_warnings.append(f"Text parsing failed: {e}")
            return "", metadata


class DocumentParserRegistry:
    """Registry for managing document parsers."""

    def __init__(self):
        self.parsers = [
            PDFParser(),
            DOCXParser(),
            MarkdownParser(),
            HTMLParser(),
            TextParser(),  # Keep as fallback
        ]

    def get_parser(self, file_path: Path) -> Optional[DocumentParser]:
        """Get the appropriate parser for a file."""
        for parser in self.parsers:
            if parser.can_parse(file_path):
                return parser
        return None

    def parse_document(self, file_path: Path) -> Tuple[str, DocumentMetadata]:
        """Parse a document using the appropriate parser."""
        logger.info(f"[DEBUG] Attempting to parse document: {file_path}")
        parser = self.get_parser(file_path)
        if parser:
            logger.info(f"[DEBUG] Using parser: {parser.__class__.__name__}")
            result = parser.parse(file_path)
            logger.info(f"[DEBUG] Parser returned content length: {len(result[0]) if result[0] else 0}")
            return result
        else:
            logger.warning(f"[DEBUG] No suitable parser found for {file_path}")
            # Fallback to text parser for unknown formats
            metadata = DocumentMetadata()
            metadata.filename = file_path.name
            metadata.file_extension = file_path.suffix.lower()
            metadata.parser_used = "Unknown"
            metadata.extraction_warnings.append("No suitable parser found")
            return "", metadata


# Global parser registry instance
document_parser_registry = DocumentParserRegistry()